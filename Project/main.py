from flask import Flask, render_template, request, send_file
from docx import Document
import requests
from bs4 import BeautifulSoup
import googletrans  # Google Translate API
import os

app = Flask(__name__)

# Function to extract text from a webpage
def extract_text_from_url(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        return " ".join([p.text for p in soup.find_all("p")])  # Extracting all <p> text
    except Exception as e:
        return f"Error fetching text: {str(e)}"

# Function to translate text
def translate_text(text, target_language):
    translator = googletrans.Translator()
    return translator.translate(text, dest=target_language).text

# Function to create a Word document
def create_word_doc(original, summarized, translated):
    doc = Document()
    doc.add_heading("Translated Content Report", level=1)

    doc.add_heading("Original Content", level=2)
    doc.add_paragraph(original)

    if summarized:
        doc.add_heading("Summarized Content", level=2)
        doc.add_paragraph(summarized)

    doc.add_heading("Translated Content", level=2)
    doc.add_paragraph(translated)

    file_path = "static/translated_output.docx"
    doc.save(file_path)
    return file_path

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/process", methods=["POST"])
def process():
    url = request.form["url"]
    language_option = request.form["language"]
    option = request.form["option"]

    # Map language option to Google Translate codes
    language_map = {
        "1": "ta", "2": "hi", "3": "te", "4": "or", "5": "kn",
        "6": "ml", "7": "bn", "8": "gu", "9": "mr", "10": "pa",
        "11": "ur", "12": "en"
    }

    target_language = language_map.get(language_option, "en")
    
    # Extract text from the URL
    original_text = extract_text_from_url(url)
    
    # Summarization (basic example - first 3 sentences)
    summarized_text = " ".join(original_text.split(".")[:3]) if option == "2" else None

    # Translate original or summarized text
    text_to_translate = summarized_text if summarized_text else original_text
    translated_text = translate_text(text_to_translate, target_language)

    # Generate Word document
    doc_path = create_word_doc(original_text, summarized_text, translated_text)

    return render_template("result.html", original=original_text, summarized=summarized_text, translated=translated_text, doc_path=doc_path)

@app.route("/download")
def download():
    return send_file("static/translated_output.docx", as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)